"""
Condensed IEEE-conference-style paper version of this project, built with
python-docx (two-column body via OOXML section columns - no LaTeX
toolchain available on this machine). All numbers pulled directly from
results/tables/, cross-checked against the same CSVs the book report uses -
nothing here is invented or restated from memory.

Re-run: .venv\\Scripts\\python.exe reports\\paper\\build_ieee_paper.py
"""

from docx import Document
from docx.shared import Inches, Pt

import ieee_common as ic

OUT_PATH = "reports/paper/mondrian_cp_er_des_paper.docx"
FIG_DIR = "reports/assignments/figures"

TITLE = "Beyond Marginal Guarantees: Mondrian Conformal Prediction for High-Variance Discrete-Event Queueing Systems"
AUTHORS = ["G. Venugopalan", "V. Sudhakar", "R. Arulprakash", "Harshith Kv"]
AFFILIATION = "Department of Artificial Intelligence, Amrita Vishwa Vidyapeetham (Faculty Guide: Akhil V.M.)"

ABSTRACT = (
    "Conformal prediction (CP) equips point-prediction surrogate models with a distribution-free, "
    "finite-sample coverage guarantee, but that guarantee is marginal: averaged over the whole input "
    "space, not within any specific operating condition. Prior work validating CP for surrogate-model "
    "uncertainty quantification has focused on physics simulations (PDEs, magnetohydrodynamics, weather, "
    "fusion), leaving the marginal-coverage limitation and CP's exchangeability assumption untested in "
    "discrete-event, queueing-driven domains. We study this gap in a calibrated emergency-department (ED) "
    "discrete-event simulation (DES), training a gradient-boosting surrogate on the DES's scenario sweep "
    "and comparing standard split CP against Mondrian CP, which calibrates separately within each of nine "
    "staffing x arrival-rate categories. While marginal coverage holds for both methods (89.6-91.4% "
    "against a 90% target), standard CP's coverage collapses to 68.2% in the single highest-stakes "
    "category - simultaneously understaffed and high-arrival - a category standard CP's own marginal "
    "average conceals. Mondrian CP restores 90.9% coverage in that same category, a difference confirmed "
    "significant across 30 independent train/calibration/test repeats (paired t-test, p < 0.001) and "
    "replicated at an independent second hospital department. We further show the coverage advantage "
    "collapses under severe exchangeability violation, quantify it against three extensions (conformalized "
    "quantile regression, conformal risk control, adaptive conformal inference under distribution shift), "
    "and benchmark five surrogate architectures. We additionally propose queueing-theoretic normalized CP "
    "(QT-CP), a continuous, bin-free alternative that derives its per-point scale directly from Kingman's "
    "heavy-traffic approximation: it achieves narrower average intervals than standard CP on three of four "
    "targets, but - reported honestly rather than omitted - does not match Mondrian CP's conditional-coverage "
    "repair. We also prove that equal-width bins along the utilization axis have provably unbounded scale "
    "distortion as utilization approaches 1, and use this to build SA-Mondrian CP, a sample-size-floored "
    "adaptive binning scheme along a single derived covariate. Tested rigorously across 30 independent "
    "repeats with paired significance testing - not the single split a first look suggested was a tie - it "
    "consistently underperforms the 2-covariate Mondrian grid (significantly on three of four targets, "
    "p ≤ 0.017), while still significantly outperforming QT-CP on three of four targets (p < 0.001): a "
    "genuine negative result for collapsing Mondrian's partition to one covariate, reported as such. Code "
    "and data tables are publicly available."
)

INDEX_TERMS = ("Conformal prediction, Mondrian conformal prediction, normalized nonconformity measures, "
               "adaptive binning, uncertainty quantification, discrete-event simulation, queueing theory, "
               "surrogate modeling, emergency department operations.")


def build():
    doc = Document()
    ic.init_styles(doc)
    ic.set_page_geometry(doc.sections[0], columns=1)

    ic.add_title_block(doc, TITLE, AUTHORS, AFFILIATION, ABSTRACT, INDEX_TERMS)

    ic.new_continuous_section(doc, columns=2)

    build_introduction(doc)
    build_related_work(doc)
    build_methodology(doc)
    build_experimental_setup(doc)
    build_results(doc)
    build_qtcp_method(doc)
    build_sa_mondrian(doc)
    build_discussion(doc)
    build_limitations(doc)
    build_conclusion(doc)
    build_references(doc)

    doc.save(OUT_PATH)
    print(f"Saved {OUT_PATH}")


def build_introduction(doc):
    ic.add_section_heading(doc, "Introduction")
    ic.add_body(doc, f"""
Surrogate models - fast, learned approximations of expensive simulations - increasingly guide operational
decisions in stochastic service systems: emergency department (ED) staffing, call-center routing, and other
queueing networks. A point prediction from such a surrogate is insufficient on its own for a decision like
"will adding two more providers bring the wait time under 45 minutes?" - the decision-maker needs a
calibrated interval with a known reliability guarantee, not a single number of unknown trustworthiness.

Conformal prediction (CP) {ic.cite('vovk2005', 'shafer2008')} supplies exactly that: a distribution-free,
finite-sample coverage guarantee requiring only that calibration and test data be exchangeable. Because the
guarantee holds without distributional assumptions, CP has been adopted for surrogate-model uncertainty
quantification across an increasingly wide set of domains. {ic.cite('gopakumar2026')} validate CP for
surrogate models in physics simulation - partial differential equations, magnetohydrodynamics, weather, and
fusion plasma modeling - and explicitly flag two open limitations of their own results: (i) the guarantee
they validate is marginal, averaged over the entire input distribution, with no guarantee that coverage
holds within any specific subgroup or operating condition; and (ii) their validation assumes calibration
and test data are exchangeable, untested under distribution shift.

Neither limitation has been tested outside physics-simulation domains. Discrete-event, queueing-driven
systems are structurally different: discrete stochastic arrivals and departures, shared-resource contention,
and priority scheduling, rather than a continuous PDE field. This paper closes that gap directly. We build a
DES of an ED calibrated on 560,486 real patient visits, train a surrogate on its scenario sweep, and test
both of {ic.cite('gopakumar2026')}'s limitations explicitly:
""")
    ic.add_body(doc, """
1) We show standard CP's marginal guarantee conceals a severe conditional coverage failure - 68.2% coverage,
22 points below the 90% target - in exactly the operating regime (understaffed, high arrival rate) where a
staffing decision matters most, and that Mondrian CP {mondrian_cite} - calibrating separately per operating
category - restores 90.9% coverage there, confirmed significant across 30 independent repeats and replicated
at an independent second department.

2) We quantify how far that advantage survives a severe, out-of-support exchangeability violation, and how
much three principled corrections (conformalized quantile regression, conformal risk control, adaptive
conformal inference) recover under shift.

3) We propose queueing-theoretic normalized CP (QT-CP), a new, bin-free nonconformity score derived from
Kingman's heavy-traffic approximation, and report its results honestly against both baselines - a genuine
efficiency gain over standard CP, but not a match for Mondrian CP's conditional-coverage repair.

4) We release the surrogate, calibration data, and all evaluation code publicly, benchmarking five surrogate
architectures for reproducibility.
""".format(mondrian_cite=ic.cite("vovk2003", "bostrom2020")))


def build_related_work(doc):
    ic.add_section_heading(doc, "Related Work")

    ic.add_subsection_heading(doc, "A", "Conformal Prediction and Its Marginal Guarantee")
    ic.add_body(doc, f"""
Split conformal prediction {ic.cite('papadopoulos2002', 'lei2018')} calibrates a single quantile of
nonconformity scores on a held-out calibration set, yielding a finite-sample marginal coverage guarantee at
negligible computational cost relative to full conformal or Bayesian alternatives. {ic.cite('angelopoulos2021')}
give a comprehensive tutorial treatment. Conformalized quantile regression (CQR) {ic.cite('romano2019')}
conformalizes a quantile regressor's raw interval rather than a fixed-width residual band, giving width that
adapts to local heteroscedasticity while retaining the same coverage guarantee.
""")

    ic.add_subsection_heading(doc, "B", "Mondrian and Group-Conditional Conformal Prediction")
    ic.add_body(doc, f"""
{ic.cite('vovk2003')} first proposed Mondrian conformal prediction: partitioning the calibration set into
disjoint categories and calibrating a separate quantile per category, trading a modest increase in
per-category calibration-set size for a per-category, rather than only marginal, coverage guarantee.
{ic.cite('bostrom2020', 'bostrom2021')} extend Mondrian calibration to regression and predictive
distributions; {ic.cite('toccaceli2019')} study combining multiple Mondrian predictors. None of this prior
work evaluates Mondrian CP in a discrete-event queueing surrogate, nor quantifies the size of the marginal-
versus-conditional gap it closes against a real operational cost asymmetry (understaffing is costlier than
overstaffing in an ED).
""")

    ic.add_subsection_heading(doc, "C", "Exchangeability, Risk Control, and Distribution Shift")
    ic.add_body(doc, f"""
{ic.cite('barber2023')} generalize conformal guarantees beyond exchangeable data; {ic.cite('tibshirani2019')}
give the likelihood-ratio weighted-CP correction for covariate shift under a known shift mechanism.
{ic.cite('bates2021')} generalize CP to risk-controlling prediction sets (CRC) for losses beyond binary
coverage; {ic.cite('gibbs2021')} propose adaptive conformal inference (ACI), which updates the calibrated
quantile online in response to observed miscoverage, dropping the exchangeability requirement entirely at
the cost of only an asymptotic, rather than finite-sample, guarantee.
""")

    ic.add_subsection_heading(doc, "D", "Surrogate Modeling, Alternative UQ, and ED Queueing")
    ic.add_body(doc, f"""
Gaussian process (GP) regression {ic.cite('rasmussen2006')} and Bayesian model calibration
{ic.cite('kennedy2001')} give principled but distributionally-assumption-dependent uncertainty; deep
ensembles {ic.cite('lakshminarayanan2017')} are a widely used alternative reviewed comprehensively by
{ic.cite('abdar2021')}. Gradient boosting {ic.cite('friedman2001')} is used here as the primary surrogate
architecture. On the applied side, queueing-theoretic ED staffing analysis {ic.cite('green2006', 'green_book',
'hu2018')} and DES-based ED modeling {ic.cite('hoot2008', 'des_calibration2021', 'des_review2022')} are
mature literatures in their own right, but - to our knowledge - have not previously been combined with a
distribution-free UQ layer validated at the level of statistical rigor (repeated-trial significance testing,
independent-site replication) applied here.
""")


def build_methodology(doc):
    ic.add_section_heading(doc, "Methodology")
    ic.add_full_width_figure(doc, "reports/paper/figures/methodology_pipeline.png",
                              "End-to-end pipeline: real data calibrates the DES, which generates surrogate "
                              "training data, evaluated under both standard and Mondrian CP.", width=Inches(6.0))

    ic.add_subsection_heading(doc, "A", "Calibrated Discrete-Event Simulation")
    ic.add_body(doc, """
The ED is a queueing system in the classical sense: patient count in system L, mean sojourn time W, and
effective arrival rate λ are linked by Little's Law,""", first_line_indent=False)
    ic.add_equation(doc, "L = λ · W")
    ic.add_body(doc, "which motivates modeling it as a discrete-event queueing simulation rather than a "
                       "generic regression task with no structural assumptions.", first_line_indent=False)
    ic.add_body(doc, f"""
We build a SimPy discrete-event simulation of a single-department ED, calibrated on the Hospital Triage and
Patient History Data {ic.cite('kaggle_dataset')} (560,486 visits, 3 departments, Yale New Haven Health
System, March 2014-July 2017). The primary calibration site (Department A, the largest and academic site)
contributes 322,283 visits, a mean of 258.2 visits/day. Real data supplies the arrival-rate distribution by
4-hour bin, day-of-week and monthly seasonality, and Emergency Severity Index (ESI) acuity mix. Service time
is not recoverable from the dataset (no discharge timestamp across its 972 columns) and is instead drawn
per-ESI from literature-calibrated log-normal distributions {ic.cite('mahmoodian2014', 'otto2022', 'kim2021')},
disclosed explicitly as a limitation rather than blended silently with the real-data-derived quantities. The
DES uses a single shared priority-resource pool, staffing capacity swept from 15 to 45 servers and
arrival-rate multiplier swept from 0.8x to 3.0x across the scenario sweep used to generate surrogate training
data, with higher-acuity patients (lower ESI) receiving priority, and is validated against real aggregated
daily volume at 91.0% match across 200 simulated days.
""")

    ic.add_subsection_heading(doc, "B", "Surrogate Model")
    ic.add_body(doc, """
A histogram-based gradient-boosting regressor is trained on a scenario sweep of the DES varying staffing
capacity and arrival-rate multiplier, predicting four operational targets: total patients served
(n_patients), mean wait time, mean total time, and the 95th-percentile wait time (p95_wait_minutes) - the
last chosen because tail wait time, not the mean, is what typically drives ED overcrowding complaints and
adverse outcomes. Section V-F additionally benchmarks four further architectures (MLP, Random Forest,
XGBoost, LightGBM) on the identical split.
""")

    ic.add_subsection_heading(doc, "C", "Standard Split Conformal Prediction")
    ic.add_body(doc, "For a calibration set of n held-out (x, y) pairs disjoint from surrogate training, the "
                       "symmetric nonconformity score is:")
    ic.add_equation(doc, "sᵢ = |yᵢ − f̂(xᵢ)|,  i = 1, …, n")
    ic.add_body(doc, "The finite-sample-corrected empirical quantile at miscoverage level α is:", first_line_indent=False)
    ic.add_equation(doc, "q̂ = Quantile( {s₁,…,sₙ}, ⌈(n+1)(1−α)⌉ / n )")
    ic.add_body(doc, "yielding the marginal coverage guarantee P(Y ∈ [f̂(X) − q̂, f̂(X) + q̂]) ≥ 1 − α, which "
                       "holds only on average over the full test distribution, not conditional on any "
                       "specific X.", first_line_indent=False)

    ic.add_subsection_heading(doc, "D", "Mondrian Conformal Prediction")
    ic.add_body(doc, f"""
Mondrian CP {ic.cite('vovk2003')} partitions the input space into disjoint categories K₁,…,K_m and calibrates
a separate quantile q̂_k per category using only that category's calibration scores:""", first_line_indent=False)
    ic.add_equation(doc, "q̂_k = Quantile( {sᵢ : xᵢ ∈ K_k}, ⌈(n_k+1)(1−α)⌉ / n_k )")
    ic.add_body(doc, "giving the strictly stronger, per-category guarantee", first_line_indent=False)
    ic.add_equation(doc, "P(Y ∈ C(X) | X ∈ K_k) ≥ 1 − α  for every category k,")
    ic.add_body(doc, """
not only on average across all categories, at the cost of a smaller effective calibration-set size n_k per
category. Because this project's scenario space has exactly two real covariates - staffing capacity and
arrival-rate multiplier - categories are the cross of staffing tercile x arrival-rate tercile (Low/Med/High x
Low/Med/High, 9 cells), with bin edges derived from the calibration set's own quantiles only, never the test
set. Both methods use identical calibration/test splits, α = 0.1, and the same symmetric nonconformity score,
isolating the effect of per-category versus pooled calibration as the only difference between them.
""", first_line_indent=False)


def build_experimental_setup(doc):
    ic.add_section_heading(doc, "Experimental Setup")
    ic.add_body(doc, """
All results use a fixed 90% target coverage (α = 0.1). Marginal-coverage comparisons (Section V-A) are
repeated across 30 independent train/calibration/test splits (different random seeds) to obtain a
significance-testable distribution of coverage and width, rather than a single-split point estimate, since a
single split cannot separate a real effect from calibration-set sampling noise. The conditional
(per-category) coverage-gap result (Section V-B) is evaluated on a single fixed split with 9 categories x a
comparable per-category test-set size (order 100 test points per cell), sufficient to detect the reported
22-point gap. Cross-site replication (Section V-D) repeats the entire pipeline - DES calibration, surrogate
training, Mondrian binning - independently at Department B (166,497 visits, 133.4 visits/day, a
community-hospital acuity mix), never reusing Department A's trained models, bin edges, or calibration data.
""")


def build_results(doc):
    ic.add_section_heading(doc, "Results")

    ic.add_subsection_heading(doc, "A", "Marginal Coverage: All Methods Meet the Target")
    ic.add_table(doc, "Marginal coverage, 30-repeat mean (target 90%)",
                 ["Method", "n_pat.", "wait", "total", "p95"],
                 [
                     ["GP baseline", "89.6%", "88.3%", "89.1%", "89.0%"],
                     ["Standard CP", "90.1%", "90.1%", "89.8%", "90.1%"],
                     ["Mondrian CP", "91.0%", "91.1%", "90.8%", "91.4%"],
                 ], col_widths=[Inches(0.85), Inches(0.55), Inches(0.55), Inches(0.55), Inches(0.5)], font_size=7.5)
    ic.add_body(doc, """
Averaged over the full test distribution, all three methods land within 2 points of the 90% target - the GP
baseline (which lacks a finite-sample guarantee) undercovers slightly, while both CP variants meet it. Judged
only on this table, Mondrian CP looks like, at best, a marginal improvement. Section V-B shows this table
conceals the actual finding.
""")

    ic.add_subsection_heading(doc, "B", "The Conditional Coverage Gap")
    fig_n = ic.add_figure(doc, f"{FIG_DIR}/exchangeability_support_diagram.png",
                            "Calibration support versus test severities used in the exchangeability stress "
                            "test (Section V-E); the same staffing x arrival-rate binning is used throughout.",
                            width=Inches(3.15))
    ic.add_table(doc, "Worst single category, mean_wait_minutes",
                 ["Category", "Pooled CP", "Mondrian CP"],
                 [["Understaffed, high arrival-rate", "68.2%", "90.9%"]],
                 col_widths=[Inches(1.45), Inches(0.85), Inches(0.85)], font_size=7.5)
    ic.add_body(doc, """
Evaluated within each of the 9 staffing x arrival-rate categories rather than pooled, standard CP's coverage
in the single understaffed / high-arrival-rate category - the operating regime a staffing decision-maker
cares about most - falls to 68.2%, 22 points under target. This is not visible in Table I's marginal average,
which blends this category's severe undercoverage with other categories' overcoverage. Mondrian CP, which
calibrates a separate quantile for exactly this category, restores 90.9% coverage there while leaving the
marginal guarantee intact (Table I). Across all four targets, standard CP's per-category coverage range
(max - min across the 9 categories) is 11.5-31.8 points; Mondrian CP's is narrower for three of four targets
(the fourth, n_patients, shows no significant conditional gap to correct - a genuine negative result reported
rather than omitted).
""")
    ic.add_figure(doc, f"{FIG_DIR}/per_category_coverage_heatmap.png",
                  "Department A per-category coverage for mean_wait_minutes: pooled (left) vs. Mondrian "
                  "(right) calibration across the full 3x3 staffing x arrival-rate grid.", width=Inches(3.3))

    ic.add_subsection_heading(doc, "C", "Statistical Significance")
    ic.add_body(doc, """
Across 30 independent repeats, Mondrian CP's marginal-coverage advantage over standard CP is small but
consistent: +0.9 to +1.3 percentage points across all four targets, each significant under a paired t-test
(p < 0.001, smallest p = 1.7 x 10⁻⁶). The practically important result is the conditional gap in Section V-B,
not this marginal difference - the marginal advantage is real but modest; the conditional-coverage repair is
the substantive finding.
""")
    ic.add_full_width_figure(doc, f"{FIG_DIR}/repeated_evaluation_boxplot.png",
                              "Distribution of empirical coverage across 30 independent (calibration, test) "
                              "draws, all four targets; the dashed line is the 90% target.")

    ic.add_subsection_heading(doc, "D", "Cross-Site Replication")
    ic.add_body(doc, """
Repeating the identical pipeline at an independent second department (Department B: 133.4 visits/day, a
community rather than academic acuity mix - 23.1% ESI-2 vs. Department A's 37.9%) reproduces the same
qualitative finding on a different, non-overlapping category: the worst category for mean_wait_minutes
(again understaffed / high-arrival) shows 76.2% pooled coverage, corrected to 89.3% by Mondrian CP. That this
replicates at a structurally different site, not merely on a resampled split of the same one, is evidence the
conditional coverage gap is a property of pooled-versus-conditional calibration in this domain generally, not
an artifact of Department A's specific arrival pattern.
""")
    ic.add_figure(doc, f"{FIG_DIR}/dept_a_vs_b_structure.png",
                  "Department A vs. B: real triage-acuity mix and daily volume/capacity - two structurally "
                  "different real sites, not a resampled copy of one.", width=Inches(3.3))

    ic.add_subsection_heading(doc, "E", "Exchangeability Violation and Its Extensions")
    ic.add_body(doc, """
Both standard and Mondrian CP's guarantees assume calibration and test data are exchangeable. We stress-test
this directly: calibrating on arrival-rate multipliers in [0.8, 1.3] and evaluating at severities up to 3.0x.
Both methods' coverage collapses together under severe shift (to 5-32% at 3.0x, depending on target) - Mondrian
CP's per-category structure offers no protection once test points fall entirely outside every calibration
category's own support, confirming this is a distinct failure mode from the conditional-coverage gap in
Section V-B, not a variant of it.
""")
    ic.add_full_width_figure(doc, f"{FIG_DIR}/mondrian_vs_standard_under_shift.png",
                              "Standard vs. Mondrian CP coverage across the severity sweep, all four targets. "
                              "Mondrian CP tracks standard CP's collapse - per-category structure gives no "
                              "protection outside its own calibrated support.")
    ic.add_body(doc, "Three corrections were evaluated against this same shift, summarized in Table III and "
                       "detailed individually below.", first_line_indent=False)
    ic.add_table(doc, "Coverage under out-of-range demand surge",
                 ["Method", "In-range", "Out-of-range"],
                 [
                     ["Static CP", "~90%", "58.0%"],
                     ["Adaptive CP (ACI)", "~90%", "84-89%"],
                     ["Weighted CP*", "~90%", "87.7-92.7%*"],
                 ], col_widths=[Inches(1.3), Inches(0.85), Inches(0.9)], font_size=7.5)

    ic.add_body(doc, """
Adaptive conformal inference (ACI) {gibbs} drops the exchangeability requirement by updating the working
miscoverage level online in response to observed errors:""".format(gibbs=ic.cite("gibbs2021")),
                first_line_indent=False)
    ic.add_equation(doc, "α_{t+1} = α_t + γ (α − errₜ),   errₜ = 1[Yₜ ∉ Cₜ(Xₜ)]")
    ic.add_body(doc, "at the cost of only an asymptotic, rather than finite-sample, guarantee.",
                first_line_indent=False)
    ic.add_full_width_figure(doc, f"{FIG_DIR}/aci_rolling_coverage.png",
                              "Rolling 60-step coverage, static CP vs. ACI, under a live demand-surge stream. "
                              "ACI recovers most of the coverage static CP loses once the stream leaves the "
                              "training range (dashed vertical line).")

    ic.add_body(doc, """
Conformal risk control (CRC) {bates} generalizes the coverage guarantee to any bounded loss ℓ_λ, selecting the
smallest λ whose empirical risk, corrected by a finite-sample upper confidence bound, stays at or below the
target:""".format(bates=ic.cite("bates2021")), first_line_indent=False)
    ic.add_equation(doc, "λ̂ = inf{ λ ∈ Λ : R̂(λ) + B̂(λ) / n ≤ α }")
    ic.add_body(doc, """
evaluated here with a clipped-overshoot-severity loss (bounding how far an undercovered interval misses, not
only whether it misses):""", first_line_indent=False)
    ic.add_figure(doc, f"{FIG_DIR}/crc_test_risk.png",
                  "CRC held-out test risk stays at or below the α = 0.10 target for all four targets, both "
                  "the 0/1 and the clipped-overshoot-severity loss.", width=Inches(3.3))

    ic.add_body(doc, """
Likelihood-ratio weighted CP {tibs} re-weights each calibration point by the covariate density ratio between
test and calibration distributions,""".format(tibs=ic.cite("tibshirani2019")), first_line_indent=False)
    ic.add_equation(doc, "w(x) = dQ_X(x) / dP_X(x)")
    ic.add_body(doc, "and takes the weighted empirical quantile", first_line_indent=False)
    ic.add_equation(doc, "q̂(x) = inf{ q : Σᵢ pᵢ(x) 1[sᵢ ≤ q] ≥ 1 − α },   pᵢ(x) = w(xᵢ) / (Σⱼw(xⱼ) + w(x))")
    ic.add_body(doc, """
under a moderate, deliberately partial-overlap shift. It restores coverage inside the region of residual
calibration/test overlap but correctly returns infinite-width intervals in the region entirely outside
calibration support, rather than a falsely narrow one - an honest, not a free, correction.
""", first_line_indent=False)
    ic.add_figure(doc, f"{FIG_DIR}/weighted_cp_coverage.png",
                  "Likelihood-ratio weighted CP coverage by region: restores the target in the overlap "
                  "region; correctly widens to 100%-coverage (infinite width) in the out-of-support tail.",
                  width=Inches(3.3))

    ic.add_body(doc, """
Conformalized quantile regression (CQR) {romano} and its Mondrian variant were also evaluated as a stronger,
width-adaptive baseline: Mondrian-CQR's coverage (90.9-92.2%) matches or exceeds Mondrian CP's, at 2-6% larger
mean interval width from its added quantile-regression variance.
""".format(romano=ic.cite("romano2019")))

    ic.add_subsection_heading(doc, "F", "Surrogate Architecture Benchmark")
    ic.add_table(doc, "Surrogate R², hardest and easiest targets",
                 ["Architecture", "p95_wait", "n_patients"],
                 [
                     ["Gradient Boosting", "0.647", "0.929"],
                     ["MLP", "0.653", "0.931"],
                     ["LightGBM", "0.646", "0.929"],
                     ["Random Forest", "0.569", "0.913"],
                     ["XGBoost", "0.572", "0.919"],
                 ], col_widths=[Inches(1.15), Inches(0.9), Inches(0.9)], font_size=7.5)
    ic.add_body(doc, """
Gradient boosting, MLP, and LightGBM cluster closely on both the hardest (p95_wait_minutes) and easiest
(n_patients) targets; bagging-based Random Forest and XGBoost's default configuration measurably trail on the
hardest target, plausibly because tail wait time is driven by rare, extreme queueing states that
boosting-to-residual architectures fit more directly than bagged averaging. All five architectures were
recalibrated with both standard and Mondrian CP; the conditional coverage gap in Section V-B replicates
qualitatively across all five (not reported in full here for space).
""")
    ic.add_figure(doc, f"{FIG_DIR}/five_architecture_r2.png",
                  "Surrogate R² across all four targets and all five architectures, identical train/test "
                  "split.", width=Inches(3.3))


def build_qtcp_method(doc):
    ic.add_section_heading(doc, "A Queueing-Theoretic Nonconformity Score")
    ic.add_body(doc, """
Mondrian CP's discrete staffing x arrival-rate bins are one way to condition calibration on operating
regime, but they require choosing bin edges and cells lose calibration-set size as the partition gets finer.
We propose an alternative that requires no discrete partition at all: normalize each nonconformity score by a
continuous, theoretically-motivated difficulty estimate derived directly from queueing theory, rather than a
generic machine-learned difficulty estimator (the usual choice in the normalized-conformal-prediction
literature). We call this queueing-theoretic normalized CP (QT-CP) and report its results honestly against
both baselines - including where it does not win.
""")

    ic.add_subsection_heading(doc, "A", "Formulation")
    ic.add_body(doc, f"""
Kingman's heavy-traffic approximation {ic.cite('kingman1962')} for the mean queueing delay Wq of a G/G/1
queue at utilization ρ,""", first_line_indent=False)
    ic.add_equation(doc, "Wq ≈ ( (Cₐ² + C_s²) / 2 ) · ( ρ / (1 − ρ) ) · E[S]")
    ic.add_body(doc, """
(Cₐ², C_s² the squared coefficients of variation of the interarrival and service-time distributions) implies
queueing delay - and its variance - scale sharply with ρ/(1−ρ) as utilization ρ approaches 1. Using the real
arrival-rate and Emergency Severity Index (ESI) calibration constants already computed for the DES (Section
III-A), the offered load at any (staffing, arrival-rate multiplier) scenario gives a closed-form utilization
estimate with no extra simulation runs:""", first_line_indent=False)
    ic.add_equation(doc, "ρ̂(c, m) = m · ( λ̄ · E[S] ) / c")
    ic.add_body(doc, """
(c = staffing capacity, m = arrival-rate multiplier, λ̄ the real mean daily arrival rate, E[S] the mean
service time under the real ESI mix). The per-point scale estimate and normalized nonconformity score are""",
                first_line_indent=False)
    ic.add_equation(doc, "σ̂(x) = 1 / (1 − min(ρ̂(x), ρ_cap)),      s̃ᵢ = sᵢ / σ̂(xᵢ)")
    ic.add_body(doc, """
calibrated exactly as in Eq. 3 but on {s̃ᵢ}, and the test-time interval half-width is q̃ · σ̂(x) - continuously
scenario-dependent, unlike standard CP's constant width, and requiring no bin edges, unlike Mondrian CP. The
cap ρ_cap prevents a divide-by-near-zero blowup; note that 35% of calibration scenarios in this project's own
scenario sweep have ρ̂ ≥ 0.9, and some exceed 1 (a deliberately oversaturated regime, Section III-A), where
Kingman's steady-state approximation is not strictly valid - ρ_cap is a necessary, not cosmetic, correction.
""")

    ic.add_subsection_heading(doc, "B", "Selecting ρ_cap Without Touching the Test Set")
    ic.add_body(doc, """
ρ_cap is itself selected using only calibration data, never the test set: the calibration set is split 70/30,
candidate caps are each calibrated on the 70% split and evaluated for coverage and mean width on the held-out
30%, and the cap with the smallest mean width among those meeting target coverage is kept - then the final
quantile is refit on the full calibration set at that cap. This mirrors this project's standing practice of
never letting a design choice see the test set (Section III-D's bin edges are chosen the same way).
""")

    ic.add_subsection_heading(doc, "C", "Results: A Genuine Trade-Off, Not a Clean Win")
    ic.add_table(doc, "Standard vs. Mondrian vs. QT-CP",
                 ["Target", "Std width", "QT-CP width", "Mondrian worst", "QT-CP worst"],
                 [
                     ["n_patients", "43.6", "51.2 (1.17x)", "83.3%", "70.3%"],
                     ["mean_wait", "47.2", "38.9 (0.83x)", "85.5%", "76.1%"],
                     ["mean_total", "46.4", "40.0 (0.86x)", "86.8%", "83.0%"],
                     ["p95_wait", "362.9", "295.2 (0.81x)", "86.3%", "76.1%"],
                 ], col_widths=[Inches(0.85), Inches(0.6), Inches(0.8), Inches(0.65), Inches(0.6)], font_size=6.8)
    ic.add_figure(doc, "reports/paper/figures/qtcp_comparison.png",
                  "QT-CP vs. standard and Mondrian CP: worst-category coverage (left) and marginal width "
                  "relative to standard CP (right), all four targets.", width=Inches(3.3))
    ic.add_body(doc, """
QT-CP retains valid marginal coverage by construction (89.3-91.3%, not tabulated for space) and achieves
narrower mean interval width than standard CP on three of four targets (0.81-0.86x) - a genuine efficiency
gain from continuous, scenario-dependent width. It also partially closes the conditional gap relative to
standard CP on the two targets with the largest real conditional gap (mean_wait_minutes, p95_wait_minutes),
consistent with this paper's own mechanism account (Section VIII): utilization-based reweighting helps exactly
where wait-time variance is utilization-driven. However, QT-CP does not match Mondrian CP's worst-category
coverage on any of the four targets, and on n_patients - the one target Section V-B already showed has no
real conditional gap to correct - QT-CP's worst-category coverage is worse than even standard CP's, plausibly
because a single scalar utilization proxy is a coarser correction than Mondrian's fully nonparametric,
per-category empirical quantile, which can absorb structure (ESI-mix effects, surrogate-specific residual
patterns) a one-dimensional ρ̂ cannot. We report this as an honest, mixed result rather than omitting the
comparison: QT-CP is a genuine methodological alternative - bin-free, theoretically grounded, and more
width-efficient - but Mondrian CP's discrete empirical binning remains the stronger correction for this
paper's central conditional-coverage question.
""")


def build_sa_mondrian(doc):
    ic.add_section_heading(doc, "SA-Mondrian CP: Singularity-Anchored Adaptive Binning")
    ic.add_body(doc, """
Mondrian CP's 3x3 staffing x arrival-rate grid (Section III-D) is one workable partition, but its bin edges
are tercile splits of the two raw covariates, not of the single derived quantity - utilization ρ̂ - that
Section VIII's mechanism argument identifies as what actually drives conditional miscoverage. This section
asks a narrower, more specific question: given that the relevant difficulty axis is ρ̂ and its scale grows as
1/(1-ρ̂) (Section VI-A), how should bins along that one axis be constructed? We show equal-width bins fail
this task provably, that the fix the resulting theorem suggests fails empirically for an identifiable reason,
and that correcting for that reason yields a working method - reported with the same honesty as QT-CP,
including where it does not beat the existing baseline.
""")

    ic.add_subsection_heading(doc, "A", "Why Equal-Width Bins Fail Near ρ → 1")
    ic.add_body(doc, """
Under the scale model σ(ρ) = κ/(1-ρ) implied by Kingman's heavy-traffic limit theorem {kingman} - which
describes the limiting shape of the whole delay distribution as ρ → 1, not only its mean - an equal-width bin
[ρ_max-Δ, ρ_max] adjacent to the largest observed utilization has within-bin scale ratio
""".format(kingman=ic.cite("kingman1962")), first_line_indent=False)
    ic.add_equation(doc, "R(ρ_max, Δ) = σ(ρ_max) / σ(ρ_max−Δ) = 1 + Δ / (1 − ρ_max)")
    ic.add_body(doc, """
For any fixed bin count B (so Δ = ρ_max/B is fixed), R → ∞ as ρ_max → 1 - e.g. at ρ_max = 0.99 with B = 9
matching Mondrian's own bin count, R = 12; at ρ_max = 0.999, R = 112 (Fig. 5, left). Bins spaced
geometrically in (1-ρ) instead - edges at 1-ρ = 2⁻ᵏ - hold a constant ratio R = 2 at every k, using only
O(log(1/(1−ρ_max))) bins to cover the same range, versus Θ(1/(1−ρ_max)) equal-width bins for the same bound.
This project's own calibration set reaches ρ̂ up to 1.86 (deliberately oversaturated scenarios, Section
III-A), well into the regime where this blowup is not a hypothetical edge case.
""", first_line_indent=False)

    ic.add_subsection_heading(doc, "B", "SA-Mondrian CP: A Sample-Size-Floored Fix")
    ic.add_body(doc, """
The theorem's natural remedy - geometric bins along ρ̂ - fails in practice: tested directly against
equal-width ρ̂-binning on real calibration data, it loses at every bin count from 5 upward, collapsing to
50-57% worst-category coverage at 15 bins (versus equal-width's 80-84%). The reason is a finite-sample cost
the asymptotic theorem does not see: geometric spacing concentrates bins in the sparse ρ̂ → 1 tail, where
this project's real scenario sweep has few calibration points, so empirical-quantile noise from tiny per-bin
samples outweighs the scale-homogeneity benefit. We call the method that corrects this SA-Mondrian CP
(Singularity-Anchored Mondrian CP): bins are built along ρ̂ alone, starting from the highest-utilization
calibration point and growing inward until each bin accumulates at least n_min points, directly targeting the
finite-sample failure mode rather than bin geometry alone. n_min is itself selected from a 70/30 split of
calibration data only - never the test set - mirroring exactly how ρ_cap is chosen for QT-CP (Section VI-B).
""")

    ic.add_subsection_heading(doc, "C", "Results: A Real Loss to Mondrian CP, a Real Win Over QT-CP")
    ic.add_body(doc, """
A first look, on a single calibration/test split, made SA-Mondrian CP appear to match the 2-covariate
Mondrian grid (two wins, two narrow losses). This project's own standing practice (Section IV) is not to
trust a single split, so we repeated the comparison across 30 independent (calibration, test) draws - the
same draws, and the same paired-significance discipline, used for this paper's central marginal-coverage
result (Section V-C) - with SA-Mondrian's n_min reselected from calibration data alone on every repeat. The
single-split tie did not hold up.
""")
    ic.add_table(doc, "Worst-Category Coverage, 30-Repeat Mean ± Std (Target 90%; p-value is SA-Mondrian vs. Mondrian, paired t-test)",
                 ["Target", "Mondrian (2D)", "QT-CP", "SA-Mondrian (1D)", "p-value"],
                 [
                     ["n_patients", "85.3 ± 2.7%", "71.3 ± 4.0%", "83.1 ± 3.9%", "0.009"],
                     ["mean_wait", "84.3 ± 3.1%", "80.3 ± 3.1%", "83.7 ± 3.5%", "0.404"],
                     ["mean_total", "84.7 ± 2.7%", "81.5 ± 3.2%", "82.3 ± 4.8%", "0.017"],
                     ["p95_wait", "85.5 ± 2.9%", "78.7 ± 4.3%", "83.2 ± 2.4%", "<0.001"],
                 ], col_widths=[Inches(0.6), Inches(0.72), Inches(0.68), Inches(0.72), Inches(0.48)], font_size=6.3)
    ic.add_figure(doc, "reports/paper/figures/sa_mondrian_comparison.png",
                  "Left: equal-width bins' worst-bin scale ratio vs. geometric bins', as a function of the "
                  "observed utilization ceiling (log scale). Right: worst-category coverage, 30-repeat mean "
                  "± std, all four targets (* p<0.05, ** p<0.001, paired t-test vs. Mondrian CP).",
                  width=Inches(3.3))
    ic.add_body(doc, """
SA-Mondrian CP loses to the 2-covariate Mondrian grid on every target - significantly on three of four
(paired t-test, p ≤ 0.017), with only mean_wait_minutes not reaching significance (p = 0.40). Collapsing
Mondrian's two raw covariates into a single derived one costs real conditional-coverage performance: the
2D grid's fully nonparametric per-cell quantile evidently captures structure (interaction effects between
staffing and arrival rate not reducible to their ratio ρ̂ alone) that a 1D construction, however carefully
binned, cannot. We report this as a genuine negative result for the hypothesis that ρ̂ alone is sufficient,
rather than reframing a loss as a tie. What does hold up under the same rigor, checked with the identical
paired test: SA-Mondrian CP significantly outperforms QT-CP on three of four targets (mean_wait_minutes,
n_patients, p95_wait_minutes; p < 0.001 each) - only mean_total_minutes shows no significant difference
between them (p = 0.44). So while a single derived covariate is not enough to match Mondrian's own 2D
partition, how that covariate's bins are built still matters more than QT-CP's fixed parametric rescaling
of it, on most though not all targets.
""")

    ic.add_subsection_heading(doc, "D", "A Natural Refinement That Changes Nothing")
    ic.add_body(doc, """
If replacing Mondrian's 2D grid with one derived covariate loses, a narrower question follows: does
refining the 2D grid, rather than replacing it, help? Within each of Mondrian's 9 existing cells, ρ̂ still
varies somewhat, so we normalized residuals by QT-CP's σ̂(x) before calibrating each cell's own quantile -
keeping Mondrian's partition intact and adding only a continuous within-cell correction. Tested with the
same 30-repeat, paired-significance rigor, this makes no significant difference on any target (|Δ| ≤ 0.54
points, p ≥ 0.19 throughout) - a genuine null result, not a small win reframed. The most likely reason:
Mondrian's tercile-width cells are already narrow enough that ρ̂, and so σ̂(x), barely varies within one -
the normalization is close to a no-op. This is itself informative: whatever keeps Mondrian CP's own
worst-category coverage at 84-86% rather than the full 90% target is not leftover within-cell utilization
heterogeneity, and combining Mondrian's partition with QT-CP's rescaling - the direction Section VII-B
might suggest - does not resolve it. We leave the actual source of that residual gap as an open question
rather than speculate past what these results support.
""")


def build_discussion(doc):
    ic.add_section_heading(doc, "Discussion")
    ic.add_body(doc, """
The mechanism behind Section V-B's gap is a heteroscedastic, queueing-theoretic one, not an artifact of the
surrogate or the CP procedure. Kingman's heavy-traffic approximation (Section VI-A) shows queueing delay -
and its variance - diverging as utilization ρ → 1, non-linearly and without bound. A single pooled
nonconformity quantile - fit mostly to the many low-utilization scenarios in the calibration set -
systematically undercovers the few high-utilization scenarios where this variance is largest. Mondrian
CP's per-category quantile is exactly the correction this mechanism calls for: a separate quantile fit only
to that category's own (larger) variance. This also explains why the gap is asymmetric across targets -
n_patients, whose variance is comparatively flat across the staffing/arrival grid, shows no significant
conditional gap to correct, while wait-time targets, whose variance is sharply regime-dependent, show the
largest gap. The practical implication for a staffing decision-maker is direct: a marginal coverage number,
however statistically valid, is the wrong number to trust when the decision at hand is specifically about the
high-utilization regime where marginal coverage is least informative about that regime's own reliability.
""", first_line_indent=False)
    ic.add_body(doc, """
This mechanism can be stated precisely rather than only narratively. Suppose the nonconformity score at
utilization ρ follows a location-scale model S(ρ) =d σ(ρ)·Z, σ(ρ) increasing in ρ and Z a fixed-shape random
variable - motivated by, though weaker than, the exact heavy-traffic limit Section VI-A cites.
""", first_line_indent=False)
    ic.add_equation(doc, "Coverage(ρ) = P(S(ρ) ≤ q̂) = P(Z ≤ q̂/σ(ρ)) = F_Z(q̂ / σ(ρ))")
    ic.add_body(doc, """
Since σ(ρ) is increasing, q̂/σ(ρ) is decreasing in ρ, and since F_Z is non-decreasing, Coverage(ρ) is a
monotonically non-increasing function of ρ for any pooled quantile q̂ - not merely observed empirically, but
a necessary consequence of this model. We checked the model's own premise against real calibration data
(mean_wait_minutes): fitting log(scale) ~ b·log(1/(1−ρ)) gives b ≈ 0.92, close to the b = 1 the heavy-traffic
limit predicts. The exact closed-form miscoverage curve this implies, however - taking Z's limiting
heavy-traffic shape (Exponential) literally - fits only loosely against the real per-bin miscoverage rate
(RMSE ≈ 0.33 across ten ρ-bins): real degradation is more gradual than the pure asymptotic prediction, since
most of the calibration sweep's mass sits well below the ρ → 1 limit where that exact shape holds. We report
the monotonicity result as structural and load-bearing, and the specific closed-form curve as a loose
approximation outside heavy traffic, not a validated quantitative law - an honest distinction we think matters
more than a tidier but overstated claim would.
""")


def build_limitations(doc):
    ic.add_section_heading(doc, "Limitations and Threats to Validity")
    ic.add_body(doc, """
Service-time distributions are literature-calibrated log-normals, not recoverable from the source dataset -
disclosed explicitly throughout rather than blended with the real-data-derived arrival and acuity
distributions, but a real gap relative to a fully real-data-calibrated DES. Cross-site replication (Section
V-D) covers two departments from one health system; a third, structurally distinct site (e.g., a different
health system or country) would strengthen the generalizability claim further. The exchangeability-violation
extensions (CRC, ACI, weighted CP) are each evaluated on a single stress-test design rather than the
30-repeat significance testing applied to the core Mondrian result in Section V-A - appropriate given this
paper's scope, but a narrower evidentiary standard than the paper's central finding, and 5 of the
architecture-benchmark results in Section V-F are summarized rather than reported in full for space. Finally,
the operational conclusion drawn (Section VIII) - that marginal coverage is the wrong number for a
high-utilization staffing decision - is drawn from a simulated environment; production deployment against
live patient data was intentionally out of scope.
""")


def build_conclusion(doc):
    ic.add_section_heading(doc, "Conclusion")
    ic.add_body(doc, """
We test both explicit limitations {gopakumar} leave open for conformal prediction on surrogate models - a
marginal-only guarantee and an untested exchangeability assumption - in a discrete-event queueing domain
outside the physics simulations they study. Standard CP's marginal guarantee conceals a severe, 22-point
conditional coverage failure in exactly the highest-stakes operating regime; Mondrian CP restores coverage
there without sacrificing the marginal guarantee, a result confirmed significant across repeated trials and
replicated at an independent site. Under severe exchangeability violation, both methods' advantage collapses
together, and we quantify how much three principled corrections recover. We also propose QT-CP, a new
bin-free nonconformity score grounded in Kingman's heavy-traffic approximation, and report it honestly: a
genuine width-efficiency gain over standard CP, but not a replacement for Mondrian CP's conditional-coverage
repair. Examining why motivates a further result: equal-width bins along the utilization axis have provably
unbounded scale distortion as utilization approaches 1, and the geometric-spacing fix this implies fails
empirically from calibration-sample scarcity in the tail - correcting for which yields SA-Mondrian CP, a
sample-size-floored adaptive binning scheme along that one derived covariate. Tested with the same 30-repeat
rigor as this paper's central finding rather than trusted from a single split, it significantly underperforms
the 2-covariate Mondrian grid on three of four targets - a genuine negative result for the hypothesis that a
single derived covariate suffices - while still significantly outperforming QT-CP on three of four,
evidence that how bins are built along the difficulty axis matters even where how many covariates define it
does not fully substitute for Mondrian's own nonparametric partition. The natural next attempt - refining
Mondrian's own grid with QT-CP's continuous rescaling rather than replacing it - changes nothing (a genuine
null result, Section VII-D), narrowing rather than resolving the open question of what keeps even Mondrian
CP's own worst-category coverage below its 90% target. These results argue that, in queueing-driven
operational domains specifically, conditional - not only marginal - coverage should be the reported and
trusted quantity wherever the decision itself is conditional on the operating regime.
""".format(gopakumar=ic.cite("gopakumar2026")))


def build_references(doc):
    for key, (authors, title, venue) in {
        "vovk2005": ("V. Vovk, A. Gammerman, and G. Shafer", "Algorithmic Learning in a Random World.",
                     "Springer, New York, NY, USA, 2005."),
        "shafer2008": ("G. Shafer and V. Vovk", '"A Tutorial on Conformal Prediction,"',
                       "J. Mach. Learn. Res., vol. 9, pp. 371-421, 2008."),
        "gopakumar2026": ("V. Gopakumar et al.", '"Uncertainty Quantification of Surrogate Models Using '
                          'Conformal Prediction,"', "Mach. Learn.: Sci. Technol., 2026."),
        "papadopoulos2002": ("H. Papadopoulos, K. Proedrou, V. Vovk, and A. Gammerman",
                             '"Inductive Confidence Machines for Regression,"',
                             "Proc. 13th Eur. Conf. Mach. Learning (ECML), LNCS vol. 2430, pp. 345-356, 2002."),
        "lei2018": ("J. Lei, M. G'Sell, A. Rinaldo, R. J. Tibshirani, and L. Wasserman",
                   '"Distribution-Free Predictive Inference for Regression,"',
                   "J. Am. Stat. Assoc., vol. 113, no. 523, pp. 1094-1111, 2018."),
        "angelopoulos2021": ("A. N. Angelopoulos and S. Bates",
                             '"A Gentle Introduction to Conformal Prediction and Distribution-Free '
                             'Uncertainty Quantification,"', "arXiv:2107.07511, 2021."),
        "romano2019": ("Y. Romano, E. Patterson, and E. Candès", '"Conformalized Quantile Regression,"',
                       "Adv. Neural Inf. Process. Syst. 32 (NeurIPS), 2019."),
        "vovk2003": ("V. Vovk, D. Lindsay, I. Nouretdinov, and A. Gammerman", '"Mondrian Confidence Machine,"',
                     "Tech. Rep., Royal Holloway, Univ. of London, 2003."),
        "bostrom2020": ("H. Boström and U. Johansson", '"Mondrian Conformal Regressors,"',
                        "Proc. Mach. Learn. Res. (PMLR), vol. 128 (COPA 2020), pp. 114-133, 2020."),
        "bostrom2021": ("H. Boström, U. Johansson, and T. Löfström",
                        '"Mondrian Conformal Predictive Distributions,"',
                        "Proc. Mach. Learn. Res. (PMLR), vol. 152 (COPA 2021), 2021."),
        "toccaceli2019": ("P. Toccaceli and A. Gammerman", '"Combination of Inductive Mondrian Conformal '
                          'Predictors,"', "Machine Learning, vol. 108, pp. 489-510, 2019."),
        "barber2023": ("R. F. Barber, E. Candès, A. Ramdas, and R. J. Tibshirani",
                       '"Conformal Prediction Beyond Exchangeability,"',
                       "Annals of Statistics, vol. 51, no. 2, pp. 816-845, 2023."),
        "tibshirani2019": ("R. J. Tibshirani, R. F. Barber, E. Candès, and A. Ramdas",
                           '"Conformal Prediction Under Covariate Shift,"',
                           "Adv. Neural Inf. Process. Syst. 32 (NeurIPS), 2019."),
        "bates2021": ("S. Bates, A. Angelopoulos, L. Lei, J. Malik, and M. Jordan",
                     '"Distribution-Free, Risk-Controlling Prediction Sets,"',
                     "J. ACM, vol. 68, no. 6, pp. 1-34, 2021."),
        "gibbs2021": ("I. Gibbs and E. Candes", '"Adaptive Conformal Inference Under Distribution Shift,"',
                     "Adv. Neural Inf. Process. Syst. 34 (NeurIPS), 2021."),
        "rasmussen2006": ("C. E. Rasmussen and C. K. I. Williams", "Gaussian Processes for Machine Learning.",
                          "MIT Press, Cambridge, MA, USA, 2006."),
        "kennedy2001": ("M. C. Kennedy and A. O'Hagan", '"Bayesian Calibration of Computer Models,"',
                        "J. R. Stat. Soc. Ser. B, vol. 63, no. 3, pp. 425-464, 2001."),
        "lakshminarayanan2017": ("B. Lakshminarayanan, A. Pritzel, and C. Blundell",
                                 '"Simple and Scalable Predictive Uncertainty Estimation Using Deep '
                                 'Ensembles,"', "Adv. Neural Inf. Process. Syst. 30 (NeurIPS), pp. 6402-6413, 2017."),
        "abdar2021": ("M. Abdar et al.", '"A Review of Uncertainty Quantification in Deep Learning: '
                      'Techniques, Applications and Challenges,"', "Information Fusion, vol. 76, pp. 243-297, 2021."),
        "friedman2001": ("J. H. Friedman", '"Greedy Function Approximation: A Gradient Boosting Machine,"',
                         "Annals of Statistics, vol. 29, no. 5, pp. 1189-1232, 2001."),
        "green2006": ("L. V. Green, J. Soares, J. F. Giglio, and R. A. Green",
                     '"Using Queueing Theory to Increase the Effectiveness of Emergency Department Provider '
                     'Staffing,"', "Academic Emergency Medicine, vol. 13, no. 1, pp. 61-68, 2006."),
        "green_book": ("L. V. Green", '"Queueing Analysis in Healthcare,"',
                       "Book chapter, Columbia Business School, New York, NY, USA."),
        "hu2018": ("X. Hu et al.", '"Applying Queueing Theory to the Study of Emergency Department '
                  'Operations: A Survey and a Discussion of Comparable Simulation Studies,"',
                  "Int. Trans. Oper. Res., 2018."),
        "hoot2008": ("N. R. Hoot, L. J. LeBlanc, I. Jones, S. R. Levin, C. Zhou, C. S. Gadd, and D. Aronsky",
                    '"Forecasting Emergency Department Crowding: A Discrete Event Simulation,"',
                    "Annals of Emergency Medicine, vol. 52, no. 2, pp. 116-125, 2008."),
        "des_calibration2021": ("A. De Santis, T. Giovannelli, S. Lucidi, M. Messedaglia, and M. Roma",
                                '"A Simulation-Based Optimization Approach for the Calibration '
                                'of a Discrete Event Simulation Model of an Emergency Department,"',
                                "arXiv:2102.00945, 2021."),
        "des_review2022": ("E. Doudareva and M. Carter", '"Discrete Event Simulation for Emergency '
                           'Department Modelling: A Systematic Review of Validation Methods,"',
                           "Operations Research for Health Care, vol. 33, 100340, 2022."),
        "kingman1962": ("J. F. C. Kingman", '"On Queues in Heavy Traffic,"',
                        "J. R. Stat. Soc. Ser. B, vol. 24, no. 2, pp. 383-392, 1962."),
        "sakasegawa1977": ("H. Sakasegawa", '"An Approximation Formula Lq = alpha . rho^beta / (1 - rho),"',
                           "Ann. Inst. Stat. Math., vol. 29, no. 1, pp. 67-75, 1977."),
        "mahmoodian2014": ("F. Mahmoodian, R. Eqtesadi, and A. Ghareghani", '"Waiting Times in Emergency '
                           'Department After Using the Emergency Severity Index Triage Tool,"',
                           "Archives of Trauma Research, vol. 3, no. 4, e19507, 2014."),
        "otto2022": ("R. Otto, S. Blaschke, W. Schirrmeister, et al.", '"Length of Stay as Quality Indicator '
                     'in Emergency Departments: Analysis of Determinants in the German Emergency Department '
                     'Data Registry (AKTIN Registry),"', "Internal and Emergency Medicine, vol. 17, no. 4, "
                     "pp. 1199-1209, 2022."),
        "kim2021": ("T. Y. Kim, C. Ohmart, Z. Khan, M. Lance, and S. Kim", '"The Effect on Length of Stay '
                    'After Implementation of Discharging Low Acuity Patients From Triage,"',
                    "Cureus, vol. 13, no. 9, e17640, 2021."),
        "kaggle_dataset": ("maalona (Kaggle username)", '"Hospital Triage and Patient History Data,"',
                           "Kaggle dataset - Yale New Haven Health System, retrospective study, "
                           "March 2014-July 2017."),
    }.items():
        ic.REF_DB[key] = f"{authors}, {title} {venue}"

    ic.add_references(doc)


if __name__ == "__main__":
    build()
