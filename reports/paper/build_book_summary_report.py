"""
Short IEEE-style summary report ABOUT the 237-page book
(reports/assignments/assignment1_mondrian_cp_coverage_gap_book.pdf) - a
navigational guide to what each chapter covers and where the central result
is proven, not a restatement of the book itself. Built with the same
ieee_common.py toolchain as the companion conference paper for a consistent,
professional look. Every chapter title, section number, and page number
below was read directly off the book's own rendered Table of Contents
(assignment1_mondrian_cp_coverage_gap_book.pdf, 237 pages) - nothing here is
paraphrased from memory or invented.

Re-run: .venv\\Scripts\\python.exe reports\\paper\\build_book_summary_report.py
"""

from docx import Document
from docx.shared import Inches, Pt

import ieee_common as ic

OUT_PATH = "reports/paper/book_summary_report.docx"

TITLE = "Beyond Marginal Guarantees: A Summary Report of the Full Book"
AUTHORS = ["G. Venugopalan", "V. Sudhakar", "R. Arulprakash", "Harshith Kv"]
AFFILIATION = "Department of Artificial Intelligence, Amrita Vishwa Vidyapeetham (Faculty Guide: Akhil V.M.)"

ABSTRACT = (
    "This document summarizes Beyond Marginal Guarantees, a 237-page, twelve-chapter book-length report "
    "examining whether conformal prediction's marginal coverage guarantee holds within specific operating "
    "regimes of a hospital emergency department, rather than only on average. The book trains a "
    "gradient-boosting surrogate on a calibrated discrete-event simulation of ED operations and compares "
    "standard split conformal prediction against Mondrian conformal prediction across nine staffing x "
    "arrival-rate categories. Its central finding is that standard conformal prediction's marginal "
    "guarantee conceals a coverage collapse to 68.2% in the single highest-stakes operating regime - "
    "understaffed and high-arrival - which Mondrian conformal prediction corrects to 90.9-92.0%, a result "
    "confirmed significant across 30 repeated trials and replicated at an independent second hospital "
    "department. Beyond this core result, the book develops five further conformal methods, stress-tests "
    "the exchangeability assumption to destruction, benchmarks five surrogate architectures, and translates "
    "the findings into an operational staffing dashboard - each derivation, implementation detail, and "
    "honest limitation documented in full. This report exists as a navigational guide to that material: "
    "what each chapter covers, where the central result is proven, and what a reader can safely skip."
)

INDEX_TERMS = ("Book summary, conformal prediction, Mondrian conformal prediction, discrete-event "
               "simulation, queueing theory, emergency department operations, reproducibility.")


def build():
    doc = Document()
    ic.init_styles(doc)
    ic.set_page_geometry(doc.sections[0], columns=1)

    ic.add_title_block(doc, TITLE, AUTHORS, AFFILIATION, ABSTRACT, INDEX_TERMS)

    ic.new_continuous_section(doc, columns=2)

    # ------------------------------------------------------------------
    # I. Purpose of this summary
    # ------------------------------------------------------------------
    ic.add_section_heading(doc, "Purpose of This Summary")
    ic.add_body(doc, """
At the faculty guide's request, this project's original short course report was expanded into a
book-length treatment: 237 pages across twelve numbered chapters plus two appendices, formatted as a
technical textbook - decimal-numbered headings, chapter-scoped figure/table/equation numbering, and
IEEE-style bracketed citations. A condensed, 8-page conference-style paper covering the same central
finding also exists as a companion document. This report is neither of those - it is a short guide to
the book: what each chapter contains, where its central result is actually proven, and which chapters a
reader with limited time can defer or skip.
""")

    # ------------------------------------------------------------------
    # II. How the book is organized
    # ------------------------------------------------------------------
    ic.add_section_heading(doc, "How the Book Is Organized")
    ic.add_table(doc, "Book Structure and Chapter Starting Pages",
                 ["Ch.", "Title", "Page"],
                 [
                     ["1", "The ER Uncertainty Dilemma", "11"],
                     ["2", "Foundations and Shadows (Literature Review)", "21"],
                     ["3", "The Marginal Trap (Research Gap)", "47"],
                     ["4", "Architecture of Conformal Simulation (Methodology)", "55"],
                     ["5", "Software Engineering and Reproducibility", "89"],
                     ["6", "Beyond Standard CP (Five Extension Methods)", "98"],
                     ["7", "Empirical Validation and Metamodel Benchmarking", "111"],
                     ["8", "When Exchangeability Fails", "136"],
                     ["9", "Cross-Site Generalization", "144"],
                     ["10", "Translational Health Operations", "152"],
                     ["11", "Synthesis and Uncharted Horizons", "161"],
                     ["12", "References", "166"],
                 ],
                 col_widths=[Inches(0.35), Inches(2.55), Inches(0.45)], font_size=8)

    ic.add_subsection_heading(doc, "A", "Chapters 1-3: The Problem")
    ic.add_body(doc, """
Chapter 1 motivates why a point-prediction surrogate needs a calibrated interval rather than a single
number, and why Emergency Department operations - discrete, stochastic, queueing-driven - are a
structurally different test domain than the physics simulations conformal prediction was first validated
on (Section 1.4, citing Gopakumar et al. {gopakumar}). Chapter 2 is the book's literature review in full
depth (26 pages, Sections 2.1-2.8): conformal prediction foundations, Mondrian and conditional-coverage
methods {mondrian}, surrogate modeling and UQ, queueing-theoretic ED operations research, and
discrete-event ED simulation, closing with a critical assessment (2.7) of where this project sits relative
to prior work. Chapter 3 states the research gap directly and formalizes the operational-risk framing
(3.6) that motivates treating conditional, not only marginal, coverage as the quantity that matters.
""".format(gopakumar=ic.cite("gopakumar2026"), mondrian=ic.cite("vovk2003")))

    ic.add_subsection_heading(doc, "B", "Chapters 4-5: Methodology and Implementation")
    ic.add_body(doc, """
Chapter 4 derives every theoretical result later chapters rely on: the discrete-event simulation design,
offered-load capacity sizing (4.2.1), all five uncertainty-quantification methods with formal algorithm
statements (4.4.5), and the heavy-traffic and quantile-monotonicity proofs behind them (4.7). Chapter 5
documents the software itself - dataset, module-by-module code walkthrough, and reproducibility practice
- matched by the full source-code listings in Appendix A.
""")

    ic.add_subsection_heading(doc, "C", "Chapter 6: Five Extensions to Standard CP")
    ic.add_body(doc, """
Conformalized quantile regression and Mondrian-CQR (6.1), conformal risk control for bounding operational
overflow severity rather than only its probability (6.2), adaptive conformal inference under a live
demand-surge stream (6.3), likelihood-ratio weighted CP under covariate shift (6.4), and a new
queueing-theoretic normalized CP - QT-CP - grounded in Kingman's heavy-traffic approximation {kingman}
(6.5).
""".format(kingman=ic.cite("kingman1962")))

    ic.add_subsection_heading(doc, "D", "Chapters 7-9: Results")
    ic.add_body(doc, """
Chapter 7 is the book's central empirical chapter: the coverage collapse and its correction (7.5), full
per-category breakdowns for all three affected targets (7.5.2-7.5.3), the mechanism behind why the
understaffed/high-demand category specifically fails (7.5.4), and 30-repeat statistical significance
testing (7.6). Chapter 8 stress-tests the exchangeability assumption every method here depends on, to the
point of collapse. Chapter 9 replicates the entire pipeline - simulation calibration, surrogate training,
Mondrian binning - independently at a second hospital department.
""")

    ic.add_subsection_heading(doc, "E", "Chapters 10-11: Application and Synthesis")
    ic.add_body(doc, """
Chapter 10 translates the statistical result into an operational prediction-interval dashboard and a
capacity-planning optimization built on this project's own conformal intervals. Chapter 11 synthesizes the
contributions, states open theoretical questions, and lays out a roadmap for future work, including
combining QT-CP's continuous scale with Mondrian's discrete binning.
""")

    # ------------------------------------------------------------------
    # III. The central finding
    # ------------------------------------------------------------------
    ic.add_section_heading(doc, "The Central Finding")
    ic.add_body(doc, """
Averaged across the full test distribution, standard and Mondrian CP both meet the 90% marginal coverage
target (Chapter 7.4-7.5). Evaluated per category instead of pooled, standard CP's coverage in the single
understaffed / high-arrival-rate category collapses to 68.2% - 22 points under target - a failure the
marginal average conceals entirely. Mondrian CP restores 90.9-92.0% coverage in that same category
(Table II), a difference significant under a paired t-test across 30 independent repeats (p < 0.001,
Section 7.6) and replicated at an independent second department, where the same category's coverage
improves from 76.2% to 89.3% (Chapter 9.3).
""")
    ic.add_table(doc, "Worst-Category Coverage, Understaffed / High-Arrival",
                 ["Site", "Pooled CP", "Mondrian CP"],
                 [
                     ["Dept. A (primary)", "68.2%", "90.9%"],
                     ["Dept. B (replication)", "76.2%", "89.3%"],
                 ],
                 col_widths=[Inches(1.4), Inches(0.95), Inches(0.95)], font_size=8)

    # ------------------------------------------------------------------
    # IV. What the book adds beyond the companion paper
    # ------------------------------------------------------------------
    ic.add_section_heading(doc, "What the Book Adds Beyond the Companion Paper")
    ic.add_body(doc, """
The 8-page companion paper states the same central finding and the same five extension methods in
condensed form. The book's additional length is not padding: it carries the full derivations and formal
proofs behind every method (Section 4.7), a literature review roughly six times longer than the paper's
Related Work section (Chapter 2), complete software documentation with a full source-code appendix
(Chapter 5, Appendix A), per-category result tables broken out for every target rather than summarized
for space (Appendix B), a dedicated operational chapter building an actual staffing dashboard and capacity
optimizer that the paper only mentions as future work (Chapter 10), and an extended synthesis of open
questions and a concrete future-work roadmap (Chapter 11).
""")

    # ------------------------------------------------------------------
    # V. Honest results and limitations
    # ------------------------------------------------------------------
    ic.add_section_heading(doc, "Honest Results and Limitations")
    ic.add_body(doc, """
Not every result in the book is a clean win, and it is written to say so directly. QT-CP (6.5) achieves
narrower average intervals than standard CP on three of four targets but does not match Mondrian CP's
worst-category coverage on any target, and underperforms even standard CP on n_patients - the one target
with no real conditional gap to correct in the first place (7.5.6). Both standard and Mondrian CP's
coverage advantage collapses together under severe exchangeability violation (Chapter 8): Mondrian's
per-category structure offers no protection once test points fall entirely outside every calibration
category's own support. Service-time distributions are literature-calibrated log-normals rather than
recoverable from the source dataset, and cross-site replication covers two departments from one health
system - both disclosed explicitly in Chapter 11.4 rather than left for a reader to discover.
""")

    # ------------------------------------------------------------------
    # VI. How to read the book
    # ------------------------------------------------------------------
    ic.add_section_heading(doc, "How to Read the Book")
    ic.add_body(doc, """
Readers primarily interested in the empirical findings rather than the underlying theory may proceed
directly to Chapter 7 after Chapter 3 - the book's own Preface gives this same guidance. Readers verifying
the project's engineering rather than its statistics should read Chapter 5 and Appendix A. Readers
evaluating the report for a specific operational deployment should read Chapters 9 and 10 for
generalizability and translation to practice, and Chapter 11.4 for what remains unresolved.
""")

    # ------------------------------------------------------------------
    # VII. Conclusion
    # ------------------------------------------------------------------
    ic.add_section_heading(doc, "Conclusion")
    ic.add_body(doc, """
Beyond Marginal Guarantees documents, at book length, a single central claim also stated in the companion
paper: a marginal coverage guarantee can conceal a severe conditional failure in exactly the operating
regime a decision-maker cares most about, and Mondrian conformal prediction corrects it without
sacrificing the marginal guarantee. What the extra 200-odd pages contribute is depth - full proofs, full
per-category evidence, full reproducibility - and honesty about where the project's own extensions (QT-CP)
and assumptions (exchangeability) do not fully hold up.
""")

    for key, (authors, title, venue) in {
        "gopakumar2026": ("V. Gopakumar et al.", '"Uncertainty Quantification of Surrogate Models Using '
                          'Conformal Prediction,"', "Mach. Learn.: Sci. Technol., 2026."),
        "vovk2003": ("V. Vovk, D. Lindsay, I. Nouretdinov, and A. Gammerman", '"Mondrian Confidence Machine,"',
                     "Tech. Rep., Royal Holloway, Univ. of London, 2003."),
        "kingman1962": ("J. F. C. Kingman", '"On Queues in Heavy Traffic,"',
                        "J. R. Stat. Soc. Ser. B, vol. 24, no. 2, pp. 383-392, 1962."),
    }.items():
        ic.REF_DB[key] = f"{authors}, {title} {venue}"

    ic.add_references(doc)

    doc.save(OUT_PATH)
    print(f"Saved {OUT_PATH}")


if __name__ == "__main__":
    build()
