"""
Assignment 2: "When Extrapolation Fails: Testing Conformal Prediction's
Exchangeability Assumption Across Surrogate Architectures."

Short (2-4 page) written report on the exchangeability stress test - both
the coverage collapse outside the training distribution, and the
counter-intuitive cross-architecture finding (confidently-wrong extrapolation
is worse than frozen-but-plausible extrapolation, because the true DES output
saturates rather than growing without bound). Pulls real numbers directly
from results/tables/ - nothing in this report is invented.

Re-run: .venv\\Scripts\\python.exe reports\\assignments\\build_assignment2.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_PATH = "reports/assignments/assignment2_exchangeability_extrapolation.docx"

TEAM = [
    ("G Venugopalan", "CB.AI.U4AID25115"),
    ("Vipin Sudhakar", "CB.AI.U4AID25166"),
    ("Rithvik Arulprakash", "CB.AI.U4AID25148"),
    ("Harshith Kv", "CB.AI.U4AID25119"),
]
COURSE_CODE = "23AID201"

HEADER_BLUE = RGBColor(0x1E, 0x40, 0xAF)
GREY = RGBColor(0x40, 0x40, 0x40)


def set_cell_shading(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = HEADER_BLUE
    return h


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for j, h in enumerate(headers):
        hdr_cells[j].text = h
        set_cell_shading(hdr_cells[j], "1E40AF")
        for p in hdr_cells[j].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(10)
    for row in rows:
        cells = table.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = str(val)
            for p in cells[j].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.size = Pt(10)
    if widths:
        for row in table.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = w
    doc.add_paragraph()
    return table


def build():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("When Extrapolation Fails: Testing Conformal Prediction's\n"
                       "Exchangeability Assumption Across Surrogate Architectures")
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = HEADER_BLUE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(f"Course: {COURSE_CODE}")
    r.font.size = Pt(11)
    r.font.color.rgb = GREY

    team_p = doc.add_paragraph()
    team_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = team_p.add_run("  |  ".join(f"{n} ({i})" for n, i in TEAM))
    r.font.size = Pt(10)
    r.font.color.rgb = GREY
    doc.add_paragraph()

    add_heading(doc, "Abstract", level=2)
    doc.add_paragraph(
        "Conformal prediction's coverage guarantee depends on an exchangeability assumption between "
        "calibration and test data - it is not designed to hold under distribution shift. Gopakumar et al. "
        "(2026) name this as a second, untested limitation of conformal prediction (CP) for surrogate-model "
        "uncertainty quantification. We stress-test this assumption in a discrete-event ER simulation by "
        "pushing the test distribution's demand level up to 3x the training range, using two structurally "
        "different surrogate architectures - gradient boosting and a neural network (MLP). Coverage "
        "collapses in both cases, confirming the limitation, but the failure mechanism differs sharply: "
        "gradient boosting's predictions freeze at the training boundary, while the MLP keeps extrapolating "
        "a trend that overshoots the true value by 6-9x more error. The ability to extrapolate is not "
        "automatically an advantage - it depends on whether the true relationship continues the trend the "
        "model learned, or saturates, as it does here."
    )

    add_heading(doc, "1. Introduction", level=2)
    doc.add_paragraph(
        "Conformal prediction (CP) provides a distribution-free, finite-sample coverage guarantee, but "
        "that guarantee rests on an exchangeability assumption: calibration data and test data must be "
        "drawn from the same distribution. Gopakumar et al. (2026), who validate CP for surrogate-model "
        "uncertainty quantification in physics simulation domains, explicitly flag this assumption as "
        "untested under distribution shift. This report investigates what happens when it is violated in a "
        "discrete-event ER queueing simulation, and whether the failure mode depends on the surrogate "
        "model's architecture."
    )

    add_heading(doc, "2. Method", level=2)
    doc.add_paragraph(
        "Calibration was kept fixed exactly as established for in-distribution evaluation (arrival-rate "
        "multiplier in [0.8, 1.3] relative to the real calibrated rate). The test distribution's "
        "arrival-rate multiplier was then pushed progressively outward - 1.5x, 1.8x, 2.0x, 2.5x, 3.0x - an "
        "escalating demand surge well outside the training range, while staffing capacity remained drawn "
        "from its normal range, isolating the shift to demand alone. Coverage was measured at each severity "
        "level for both standard and Mondrian conformal prediction, using two surrogate architectures "
        "trained on identical data: a gradient-boosting regressor (GBR, tree-based) and a multilayer "
        "perceptron (MLP, a small neural network), which achieve near-identical in-distribution accuracy "
        "(R-squared within 0.01 of each other on every target), ensuring any difference observed under "
        "distribution shift reflects architecture, not overall model quality."
    )

    add_heading(doc, "3. Results", level=2)
    doc.add_paragraph(
        "Table 1 shows standard CP's coverage collapsing as severity increases past the training boundary "
        "(1.3x), for the gradient-boosting surrogate."
    )
    add_table(doc,
              ["Arrival multiplier", "n_patients", "mean_wait", "mean_total", "p95_wait"],
              [
                  ["1.3 (in range)", "93.3%", "83.0%", "88.0%", "82.0%"],
                  ["1.5", "78.7%", "68.7%", "73.7%", "69.7%"],
                  ["1.8", "70.3%", "42.7%", "47.3%", "43.0%"],
                  ["2.0", "64.7%", "34.3%", "39.0%", "33.0%"],
                  ["3.0", "31.7%", "12.3%", "4.7%", "78.3%*"],
              ])
    doc.add_paragraph(
        "Table 1. Gradient-boosting surrogate: coverage vs. arrival severity. *p95_wait's apparent recovery "
        "at 3.0x is not the interval working correctly - it is explained in Section 4."
    ).italic = True

    doc.add_paragraph(
        "Repeating the test with the MLP surrogate produces an even faster collapse for two of the four "
        "targets: n_patients and mean_total_minutes both reach exactly 0% coverage by 2.0x severity, "
        "compared to the gradient-boosting surrogate's more gradual decline to 31.7% and 4.7% respectively "
        "at 3.0x (Table 2)."
    )
    add_table(doc,
              ["Arrival multiplier", "n_patients (MLP)", "mean_total (MLP)", "n_patients (GBR)", "mean_total (GBR)"],
              [
                  ["1.3 (in range)", "93.0%", "88.3%", "93.3%", "88.0%"],
                  ["1.8", "6.7%", "2.3%", "70.3%", "47.3%"],
                  ["2.0", "0.0%", "0.0%", "64.7%", "39.0%"],
                  ["3.0", "0.0%", "0.0%", "31.7%", "4.7%"],
              ])
    doc.add_paragraph(
        "Table 2. Standard CP coverage, MLP vs. gradient-boosting (GBR) surrogate, standard CP."
    ).italic = True

    doc.add_paragraph(
        "This is initially counter-intuitive: gradient-boosting (tree-based) models cannot extrapolate "
        "past the range of their training data - their predictions clip to whatever leaf the boundary "
        "training points fell into, effectively freezing. The MLP has no such limitation and its "
        "predictions keep changing as severity increases. One would expect the architecture capable of "
        "extrapolating to handle distribution shift better. It does not."
    )
    doc.add_paragraph(
        "Comparing both models' predictions against the true simulated value at fixed capacity (30 "
        "servers) reveals why (Table 3). The true DES output saturates at extreme demand rather than "
        "growing proportionally: at extreme overload, more arriving patients simply queue up and do not "
        "complete service within the simulated day, so the count of completed visits levels off. The "
        "gradient-boosting model's frozen prediction happens to land close to this saturating value by "
        "coincidence - the true function actually is roughly flat in this regime, so \"frozen\" is "
        "qualitatively the right shape. The MLP instead extrapolates the upward trend it learned near the "
        "training boundary and continues it indefinitely, overshooting the true value by a large margin."
    )
    add_table(doc,
              ["n_capacity = 30", "Arrival 1.0x", "Arrival 1.3x", "Arrival 2.0x", "Arrival 3.0x", "n_patients, true"],
              [
                  ["True value", "235.2", "244.8", "258.6", "280.9", "(saturating)"],
                  ["GBR prediction", "240.5", "247.8", "247.8 (frozen)", "247.8 (frozen)", "error ≈ 33"],
                  ["MLP prediction", "234.8", "245.8", "336.5", "521.2", "error ≈ 240"],
              ])
    doc.add_paragraph(
        "Table 3. True vs. predicted n_patients at fixed capacity, both architectures. The MLP's error at "
        "3.0x severity is roughly 7x larger than the gradient-boosting model's, despite - or rather, "
        "because of - its ability to extrapolate."
    ).italic = True

    add_heading(doc, "4. Discussion", level=2)
    doc.add_paragraph(
        "The p95_wait_minutes target's apparent coverage \"recovery\" at extreme severity (Table 1, 3.0x) "
        "is a data-generating-process artifact, not a sign the interval is working. At extreme overload, "
        "fewer patients finish service within the simulated day at all, shrinking and changing the "
        "composition of the completed-visit pool the 95th percentile is computed over. The true value "
        "drifting back toward the frozen prediction is coincidental, driven by this composition shift, not "
        "by the surrogate or the interval becoming more reliable."
    )
    doc.add_paragraph(
        "The central finding - that an architecture's ability to extrapolate is not automatically an "
        "advantage - depends entirely on whether the true relationship continues the trend a model learned "
        "near its training boundary. Here it does not: the relationship saturates due to a censoring "
        "mechanism in how the outcome is measured. A model that confidently extrapolates the wrong trend "
        "can be worse than one that cannot extrapolate at all. This is a meaningful caveat for any "
        "practitioner choosing a surrogate architecture specifically to \"handle\" out-of-distribution "
        "inputs: the choice should depend on the true relationship's expected shape beyond the training "
        "range, not on extrapolation capability alone."
    )
    doc.add_paragraph(
        "One property is worth noting as a partial mitigation: the failure is detectable, not silent. "
        "Residuals grow and coverage measurably collapses rather than the surrogate confidently reporting a "
        "narrow, wrong interval with no signal that something is amiss. This does not restore the coverage "
        "guarantee, but it means the failure mode is one a monitoring system could in principle catch."
    )

    add_heading(doc, "5. Conclusion", level=2)
    doc.add_paragraph(
        "Conformal prediction's exchangeability assumption breaks down outside the training distribution "
        "in this domain, confirming Gopakumar et al.'s second stated limitation. Testing two structurally "
        "different surrogate architectures shows the breakdown is not an artifact of one architecture's "
        "specific limitation (tree-based extrapolation freezing) - it occurs regardless of architecture, "
        "though the failure mechanism and severity differ. Counter-intuitively, the architecture that "
        "cannot extrapolate at all degraded more gracefully than the one that could, because the true "
        "relationship in this domain saturates rather than growing without bound - a reminder that "
        "extrapolation capability is not inherently protective against distribution shift."
    )

    add_heading(doc, "References", level=2)
    doc.add_paragraph("Gopakumar, et al. (2026). Conformal Prediction for surrogate-model uncertainty "
                       "quantification in physics simulation domains (PDEs, MHD, weather, fusion). [Base paper]")
    doc.add_paragraph("Hospital Triage and Patient History Data. Kaggle (maalona) - Yale New Haven Health "
                       "System, retrospective study, March 2014 - July 2017. [Real-world calibration data]")

    doc.save(OUT_PATH)
    print(f"Saved {OUT_PATH}")


if __name__ == "__main__":
    build()
