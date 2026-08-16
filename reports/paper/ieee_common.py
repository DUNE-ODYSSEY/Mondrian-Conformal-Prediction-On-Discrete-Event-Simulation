"""
Shared formatting helpers for an IEEE-conference-style two-column paper,
built with python-docx (no LaTeX toolchain available on this machine).
Approximates the standard IEEEtran conference template: Times New Roman,
two-column body with a single-column title/abstract block, numbered
[N]-style citations, Roman-numeral section headings, "Fig. N." / "TABLE N"
captions.
"""

import re

from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Cm, Twips

BLACK = RGBColor(0x00, 0x00, 0x00)
GREY = RGBColor(0x40, 0x40, 0x40)

_section_counter = {"n": 0}
_figure_counter = {"n": 0}
_table_counter = {"n": 0}
_ref_order = []  # citation keys in first-cited order
REF_DB = {}  # key -> formatted IEEE reference string, populated by caller


# --------------------------------------------------------------------------
# Page / column setup
# --------------------------------------------------------------------------

def set_page_geometry(section, columns=1):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(1.0)
    set_columns(section, columns)


def set_columns(section, n, space=Twips(360)):
    """space is in twips (1/1440in) - w:cols' w:space attribute unit, NOT EMU."""
    sectPr = section._sectPr
    for existing in sectPr.findall(qn("w:cols")):
        sectPr.remove(existing)
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), str(n))
    cols.set(qn("w:space"), str(space.twips))
    if n > 1:
        cols.set(qn("w:equalWidth"), "1")
    sectPr.append(cols)


def new_continuous_section(doc, columns):
    section = doc.add_section(WD_SECTION.CONTINUOUS)
    set_page_geometry(section, columns=columns)
    return section


# --------------------------------------------------------------------------
# Base styles
# --------------------------------------------------------------------------

def init_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0


# --------------------------------------------------------------------------
# Title block (single-column section)
# --------------------------------------------------------------------------

def add_title_block(doc, title, authors, affiliation, abstract, index_terms):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.name = "Times New Roman"
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(", ".join(authors))
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(affiliation)
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.name = "Times New Roman"
    p.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Abstract—")
    r.font.bold = True
    r.font.italic = True
    r.font.size = Pt(9)
    r2 = p.add_run(abstract)
    r2.font.italic = True
    r2.font.size = Pt(9)
    p.paragraph_format.first_line_indent = Inches(0.2)

    p = doc.add_paragraph()
    r = p.add_run("Index Terms—")
    r.font.bold = True
    r.font.italic = True
    r.font.size = Pt(9)
    r2 = p.add_run(index_terms)
    r2.font.italic = True
    r2.font.size = Pt(9)
    p.paragraph_format.first_line_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(10)


# --------------------------------------------------------------------------
# Section / body / citation helpers
# --------------------------------------------------------------------------

_ROMAN = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X", "XI", "XII"]


def add_section_heading(doc, title):
    _section_counter["n"] += 1
    n = _section_counter["n"]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{_ROMAN[n - 1]}. {title.upper()}")
    r.font.bold = True
    r.font.size = Pt(10)
    r.font.name = "Times New Roman"
    return p


def add_subsection_heading(doc, letter, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{letter}. {title}")
    r.font.bold = True
    r.font.italic = True
    r.font.size = Pt(10)
    r.font.name = "Times New Roman"
    return p


def cite(*keys):
    """Return the inline '[N]' / '[N], [M]' text and record citation order."""
    nums = []
    for k in keys:
        if k not in _ref_order:
            _ref_order.append(k)
        nums.append(_ref_order.index(k) + 1)
    return "[" + "], [".join(str(n) for n in nums) + "]"


def add_body(doc, text, justify=True, first_line_indent=True):
    for block in re.split(r"\n\s*\n", text.strip()):
        block = " ".join(line.strip() for line in block.strip().splitlines())
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        if first_line_indent:
            p.paragraph_format.first_line_indent = Inches(0.2)
        if justify:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(block)
        r.font.size = Pt(10)
        r.font.name = "Times New Roman"


def add_equation(doc, text, number=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.italic = True
    r.font.size = Pt(10)
    r.font.name = "Times New Roman"
    if number:
        r2 = p.add_run(f"\t\t({number})")
        r2.font.size = Pt(10)
        r2.font.name = "Times New Roman"
    return p


# --------------------------------------------------------------------------
# Figures and tables
# --------------------------------------------------------------------------

def add_figure(doc, path, caption, width=Inches(3.3)):
    _figure_counter["n"] += 1
    n = _figure_counter["n"]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run()
    run.add_picture(path, width=width)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(f"Fig. {n}. {caption}")
    r.font.size = Pt(8.5)
    r.font.name = "Times New Roman"
    return n


def set_cell_shading(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def add_table(doc, caption, headers, rows, col_widths=None, font_size=8):
    _table_counter["n"] += 1
    n = _table_counter["n"]
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(3)
    r = cap.add_run(f"TABLE {_ROMAN[n - 1]}\n{caption.upper()}")
    r.font.size = Pt(8.5)
    r.font.name = "Times New Roman"

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for j, h in enumerate(headers):
        hdr_cells[j].text = h
        set_cell_shading(hdr_cells[j], "D9D9D9")
        for p in hdr_cells[j].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(font_size)
                r.font.name = "Times New Roman"
    for row in rows:
        cells = table.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = str(val)
            for p in cells[j].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.size = Pt(font_size)
                    r.font.name = "Times New Roman"
    if col_widths:
        table.autofit = False
        for row in table.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = w
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return n


# --------------------------------------------------------------------------
# References
# --------------------------------------------------------------------------

def add_references(doc):
    add_section_heading(doc, "References")
    _section_counter["n"] -= 1  # References is unnumbered in IEEE style
    # Overwrite the heading we just added without a Roman numeral.
    doc.paragraphs[-1].runs[0].text = "REFERENCES"
    for i, key in enumerate(_ref_order, start=1):
        entry = REF_DB.get(key, f"[missing reference: {key}]")
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"[{i}] {entry}")
        r.font.size = Pt(8.5)
        r.font.name = "Times New Roman"
