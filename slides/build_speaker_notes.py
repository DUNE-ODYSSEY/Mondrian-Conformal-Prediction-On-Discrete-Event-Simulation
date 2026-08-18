"""
Speaker notes / talking points handout for the mid-sem PPT
(mid_sem_presentation.pptx, 14 slides). Each team member presents a
CONTIGUOUS block of slides, not an alternating round-robin - Rithvik
does the intro, then hands off to Venu for the approach/methodology
section, then Vipin for data and results, then Harshith to close out
(4/4/3/3 slides, matching the natural section breaks in the deck).
Talking points are written to match what is actually on each slide - not
a restatement of the bullets, but what a presenter would actually say
while that slide is on screen - plus a one-line handoff at the end of
each speaker's last slide.

Re-run: .venv\\Scripts\\python.exe slides\\build_speaker_notes.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DOCX = "slides/mid_sem_speaker_notes.docx"
OUT_PDF = "slides/mid_sem_speaker_notes.pdf"

# Contiguous blocks, not round-robin: each name repeated once per slide in
# their block, in slide order.
SPEAKER_BLOCKS = [
    ("Rithvik", 4),    # Slides 1-4: title, contents, problem, gap
    ("Venu", 4),       # Slides 5-8: concept diagrams, approach, methodology
    ("Vipin", 3),      # Slides 9-11: data, work completed, core finding
    ("Harshith", 3),   # Slides 12-14: DES detail, lit review, future scope
]
SPEAKERS = [name for name, count in SPEAKER_BLOCKS for _ in range(count)]

NAVY = RGBColor(0x1E, 0x40, 0xAF)
GREY = RGBColor(0x40, 0x40, 0x40)
SPEAKER_COLOR = RGBColor(0xB8, 0x3A, 0x1E)

SLIDES = [
    ("Title Slide",
     "Mondrian Conformal Prediction for Uncertainty Quantification in ER Discrete-Event Simulation Surrogates",
     [
        "Welcome everyone - we're presenting our mid-semester review for course 23AID201.",
        "Our project applies conformal prediction, specifically Mondrian conformal prediction, to a "
        "surrogate model trained on an emergency-department discrete-event simulation.",
        "Quick intros: I'm Rithvik, and with me are Venu, Vipin, and Harshith - I'll take the introduction, "
        "then each of them will take over in turn for one section apiece, so you'll hear from all four of "
        "us in sequence.",
        "We're at roughly 80 percent of the project now - what's built, validated, and checked so far, and "
        "what's left for end-sem.",
     ]),
    ("Contents",
     "Roadmap for this presentation",
     [
        "Here's the roadmap for the next few minutes.",
        "We'll start with the problem we're solving and the specific research gap we're addressing.",
        "Then two short concept slides - one explaining why a marginal coverage number can be misleading, "
        "and one showing the Mondrian partition idea - before we get into our actual bridging approach.",
        "After that: methodology, the real dataset we calibrated on, what's concretely done so far, our core "
        "finding, the simulation in more detail, a snapshot of our literature review, and finally what's "
        "left for end-sem.",
     ]),
    ("Problem Statement",
     "Why point predictions from a surrogate aren't enough",
     [
        "Surrogate models are fast, learned stand-ins for expensive simulations, and they're increasingly "
        "used to guide real decisions in systems like ER staffing or call-center queueing.",
        "But a single point prediction isn't enough for a decision like 'will adding two more doctors "
        "actually bring the wait time down' - you need a calibrated interval, not just one number.",
        "Two standard ways to get that interval: Gaussian Processes, which are theoretically solid but "
        "expensive to train and rely on distributional assumptions rather than a hard guarantee; and "
        "conformal prediction, which is distribution-free and comes with a finite-sample guarantee, but has "
        "only been validated for surrogate models in a narrow set of domains so far.",
        "That's the central question our project is built around: does conformal prediction still hold up "
        "when the surrogate is trained on a discrete-event queueing simulation, instead of the physics "
        "domains it's already been tested on?",
     ]),
    ("Research Gap",
     "What Gopakumar et al. (2026) leave open",
     [
        "Our base paper, Gopakumar et al. 2026, validates conformal prediction for surrogate-model "
        "uncertainty quantification, but only in physics domains - PDEs, magnetohydrodynamics, weather, "
        "fusion plasma modeling.",
        "They explicitly flag two open limitations of their own results: first, the guarantee they validate "
        "is marginal - averaged over the whole input space, with no guarantee it holds within any specific "
        "subgroup or condition. Second, their validation assumes calibration and test data are "
        "exchangeable, and they never test what happens under distribution shift.",
        "Neither limitation has been tested outside physics simulation. A discrete-event, queueing-driven "
        "system is structurally different - discrete stochastic arrivals and departures, shared-resource "
        "contention, priority queues - not a continuous PDE field.",
        "That's exactly the gap our project sits in, and it's the reason this domain is a genuinely new "
        "test case, not just a repeat of an already-answered question.",
        "I'll hand it over to Venu now, who'll walk through our actual approach, starting with a quick "
        "concept slide.",
     ]),
    ("Why Marginal Coverage Isn't Enough (concept)",
     "An illustrative example - not our own results",
     [
        "This chart is illustrative - a generic example, not our project's own numbers. It's here to make "
        "one specific concept concrete before we describe our approach; our own real version of this exact "
        "idea comes later, on Slide 11.",
        "Imagine four categories of a system. If you average coverage across all four, you get about 90.5 "
        "percent - it looks like your 90 percent target is comfortably met, and the red line marks exactly "
        "where that 90 percent target sits relative to each bar.",
        "But look closer: Category D alone sits at 68 percent, well below target. A decision-maker who only "
        "sees the marginal, averaged number would never know this specific category is unreliable.",
        "This is exactly the gap Mondrian conformal prediction is designed to close, by calibrating "
        "separately within each category instead of pooling everything into one marginal number.",
     ]),
    ("Our Bridging Approach",
     "Applying Mondrian CP, testing both limitations directly",
     [
        "Our core idea: apply conformal prediction, specifically Mondrian CP, to an ER queueing surrogate, "
        "and test directly whether Gopakumar et al.'s two limitations actually show up in this domain.",
        "On the left: addressing marginal coverage. Mondrian CP partitions calibration by category - "
        "staffing tercile crossed with arrival-rate tercile - instead of lumping every scenario into one "
        "marginal guarantee. That's done: Slide 11 confirms the worst category's coverage recovers from "
        "68.2 to 90.9 percent, statistically significant across 30 repeats.",
        "On the right: addressing exchangeability, with an out-of-distribution 'surge day' stress test "
        "outside the normal training range - this is the main piece of work still remaining, the core of "
        "our final ~20 percent.",
        "That stress test is where we'll get to see whether standard CP and Mondrian CP hold up, or break "
        "down, once exchangeability is genuinely violated.",
     ]),
    ("The Mondrian Partition (concept)",
     "How the 9-category grid works",
     [
        "This is the actual partition structure Mondrian CP uses - not a future plan, this is the real "
        "structure behind the result on Slide 11.",
        "Every scenario in our simulation has two covariates: staffing capacity and arrival-rate multiplier. "
        "We split each into three terciles - low, medium, high - giving a 3-by-3 grid of 9 categories.",
        "Every simulated scenario falls into exactly one of these 9 cells, based on its own staffing level "
        "and arrival-rate multiplier - nothing is shared across cells. Mondrian CP calibrates a separate "
        "quantile within each one, instead of one pooled quantile across all 9.",
        "The highlighted category is comparatively low-stress - medium staffing, low arrival rate. The "
        "opposite corner - low staffing, high arrival rate - is the real category, in our own data, where "
        "Slide 11's 68.2-to-90.9-percent recovery actually happens.",
     ]),
    ("Methodology Overview",
     "The three-stage pipeline and how each stage is checked",
     [
        "Our pipeline has three stages, shown left to right: a discrete-event simulation built in SimPy and "
        "calibrated on real ED data; a surrogate model - gradient boosting - trained on the DES's scenario "
        "sweep; and uncertainty quantification, comparing a GP baseline against standard CP and Mondrian CP.",
        "We don't just chain these stages together and hope for the best - each one is checked against the "
        "previous one before we move on.",
        "The DES's output is checked against real aggregated ED statistics. Surrogate accuracy is checked "
        "against DES outputs held out from training, using MAE, RMSE, and R-squared.",
        "And critically, all three UQ methods use the identical fixed test split and the identical target "
        "coverage of 90 percent, so GP, standard CP, and Mondrian CP end up directly comparable. All three "
        "stages are complete now - the Uncertainty Quantification stage's headline result is on Slide 11.",
        "Over to Vipin now, who'll walk through the real dataset and our actual results.",
     ]),
    ("Real-World Data Used",
     "The Kaggle dataset and what's real vs. literature-calibrated",
     [
        "We calibrate on the Hospital Triage and Patient History Data from Kaggle - real retrospective data "
        "from the Yale New Haven Health System, covering March 2014 through July 2017.",
        "It's a large dataset: 560,486 ED visits total, 972 variables per visit, across three EDs - one "
        "academic and two community sites. We calibrate on Department A only, the largest, academic site: "
        "322,283 visits, about 258 per day.",
        "We're deliberate about keeping real data and literature-sourced values separate, not blending them "
        "silently. From the real data: arrival pattern by 4-hour bin, day of week and month, ESI acuity "
        "mix, and department-level daily visit rate.",
        "From literature: service and treatment time per ESI level, modeled as log-normal - the dataset has "
        "no discharge timestamp, and we checked across all 972 columns to be sure before falling back to "
        "literature values.",
     ]),
    ("Work Completed So Far (~80%)",
     "DES validation, surrogate accuracy, and the GP baseline",
     [
        "This slide is our concrete progress checkpoint - we're at roughly 80 percent of the project now.",
        "First, DES validation: across 200 simulated days, our simulation's mean visits per day is 235.1, "
        "against a real value of 258.2 - a 91 percent match, which we consider a strong validation given how "
        "much of the calibration is real-data-driven.",
        "Second, surrogate accuracy: our gradient-boosting surrogate hits an R-squared of 0.929 on total "
        "patient count, and even the hardest target - the 95th-percentile wait time - reaches 0.647.",
        "Third, the GP baseline: at a 90 percent target, it undercovers on three of four targets, landing "
        "around 87.7 to 88.8 percent. That's not surprising - GPs rely on distributional assumptions rather "
        "than a finite-sample guarantee. Standard CP and Mondrian CP both close most of that gap on average "
        "- the next slide shows the real, more important result underneath that average.",
     ]),
    ("Core Finding: Mondrian CP Closes the Coverage Gap",
     "The real result: 68.2% -> 90.9% in the worst category",
     [
        "This is our central finding, and it's done, not projected.",
        "On average - marginal coverage - all three methods land close to the 90 percent target: GP baseline "
        "around 89 percent, standard CP and Mondrian CP both just above 90. On its own, this table would "
        "look like a minor, almost unremarkable improvement.",
        "But checked within each of the 9 categories rather than pooled, standard CP's coverage in the "
        "single worst category - understaffed with a high arrival rate - collapses to 68.2 percent, well "
        "under target. That failure is completely invisible in the marginal average above.",
        "Mondrian CP, calibrating separately per category, restores that same category to 90.9 percent "
        "coverage, without sacrificing the marginal guarantee. And this isn't a one-off lucky split - it's "
        "confirmed statistically significant across 30 independent repeats, paired t-test, p less than 0.001.",
        "I'll pass it to Harshith now to close out with the simulation in more detail, our literature "
        "review, and what's left to do.",
     ]),
    ("The Discrete-Event Simulation",
     "What the DES actually models, and how it's validated",
     [
        "This chart shows the real arrival volume by 4-hour bin that our simulation is calibrated on - you "
        "can see the clear peak in the 11-to-14 window, which matches typical ED demand patterns.",
        "Our simulation is a SimPy discrete-event model of a single ED, calibrated directly on this real "
        "arrival pattern plus the real ESI acuity mix - not a generic textbook queueing example.",
        "It uses a priority-resource pool, where higher-acuity patients - lower ESI numbers - get priority "
        "for capacity, matching how real triage actually works.",
        "And as we mentioned on Slide 10, it's validated against real aggregated daily volume at a 91 "
        "percent match across 200 simulated days - this simulation is the foundation everything downstream, "
        "the surrogate and the uncertainty quantification, is built on.",
     ]),
    ("Literature Review Snapshot",
     "20 core papers plus 10 real-dataset cross-check papers",
     [
        "Our literature review has two parts. The core review covers 20 papers across 5 categories, with 3 "
        "reviewed in depth, including Gopakumar et al. 2026, our base paper.",
        "Those 5 categories: conformal prediction foundations, Mondrian and conditional-coverage CP, "
        "surrogate modeling and uncertainty quantification, queueing theory and ED operations research, and "
        "discrete-event simulation with ED-specific machine learning.",
        "On top of that, we have 10 papers that each report real ED service-time or length-of-stay data from "
        "their own hospital datasets - we use these to cross-check our own DES's service-time calibration "
        "numerically, not just cite them for general support.",
        "The review also includes a critical assessment section - specifically looking at where the "
        "reviewed literature's own assumptions do, or don't, transfer to a discrete-event queueing domain "
        "like ours.",
     ]),
    ("Future Scope - End-Sem",
     "What's left in the remaining ~20%",
     [
        "To close: standard CP, Mondrian CP, and our core coverage-gap finding on Slide 11 are done. What's "
        "left is the remaining roughly 20 percent of the project.",
        "First, we'll stress-test exchangeability with an out-of-distribution demand-surge scenario, to see "
        "exactly where standard CP and Mondrian CP hold up, or break down, once that assumption is "
        "genuinely violated.",
        "Second, we'll finalize the full comparison - GP baseline versus standard CP versus Mondrian CP - "
        "on both coverage and interval width, across every target, not just the headline one.",
        "And finally, we'll bring it all together in the end-sem report and presentation: a fully "
        "quantified, written-up answer on whether Mondrian CP closes the marginal-versus-conditional "
        "coverage gap in this domain. Thank you.",
     ]),
]


def add_native_border(paragraph, color="1E40AF", size=8):
    pPr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    pPr.append(borders)


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Mid-Semester Review: Speaker Notes & Talking Points")
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Mondrian Conformal Prediction for Uncertainty Quantification in ER Discrete-Event "
                     "Simulation Surrogates  ·  Course 23AID201")
    r.font.size = Pt(12)
    r.font.italic = True
    r.font.color.rgb = GREY

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run("Speaking order: each person presents a contiguous block, not alternating - "
                      "Rithvik (Slides 1-4, intro) → Venu (5-8, approach & methodology) → "
                      "Vipin (9-11, data & core finding) → Harshith (12-14, detail & wrap-up)")
    r.font.size = Pt(10.5)
    r.font.color.rgb = SPEAKER_COLOR
    doc.add_paragraph()

    for i, (slide_title, slide_sub, points) in enumerate(SLIDES, start=1):
        speaker = SPEAKERS[i - 1]

        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(2)
        r = h.add_run(f"Slide {i}: {slide_title}")
        r.font.size = Pt(15)
        r.font.bold = True
        r.font.color.rgb = NAVY
        add_native_border(h)

        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(6)
        r = sp.add_run(f"Speaker: {speaker}")
        r.font.size = Pt(11)
        r.font.bold = True
        r.font.color.rgb = SPEAKER_COLOR
        r2 = sp.add_run(f"   |   On screen: {slide_sub}")
        r2.font.size = Pt(11)
        r2.font.italic = True
        r2.font.color.rgb = GREY

        for pt in points:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(pt)
            r.font.size = Pt(11.5)

    doc.save(OUT_DOCX)
    print(f"Saved {OUT_DOCX}")


if __name__ == "__main__":
    build()
